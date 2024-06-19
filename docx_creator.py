import numpy as np
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from pandas import DataFrame

from constants import *
from data.V9_402.dim1_data import dim1, str_dim1
from data.V9_402.dim2_data import dim2, str_dim2
from plot_creator import *
from stats import avg, avg_interval, delta, std, varience_interval
from table_creator import *
from sympy import N


def create_all(filename: str) -> None:
    doc = Document()
    add_paragraph(doc, "Исходные данные:",
                  align=WD_ALIGN_PARAGRAPH.CENTER, pt=16)
    add_paragraph(doc, "1. Одномерная выборка:")
    add_paragraph(doc, str_dim1, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  pt=11, font="Cambria Math")
    add_paragraph(doc, "2. Двумерная выборка:", )
    add_paragraph(doc, str_dim2, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  pt=11, font="Cambria Math")
    add_paragraph(doc, "\n1. Анализ одномерной выборки",
                  align=WD_ALIGN_PARAGRAPH.CENTER, pt=16)
    add_paragraph(doc, "1.1. Вариационный ряд\n"+variation_info)
    dim1_variation_table = variation_table(dim1)
    add_table(doc, dim1_variation_table)
    add_paragraph(
        doc, "\n1.2. Эмпирическая функция распределения\n"+distribution_info)
    add_paragraph(doc, "Расчётная таблица:")
    dim1_distribution_table = distribution_table(dim1)
    add_table(doc, dim1_distribution_table)
    add_paragraph(doc, "\nГрафик эмпирической функции распределения:")
    dim1_distribution_plot = distribution_plot(dim1)
    add_picture(doc, dim1_distribution_plot)
    add_paragraph(doc, "1.3. Равноинтервальная гистограмма относительных частот\n"+interval_hist_info)
    add_paragraph(doc, "Расчётная таблица:")
    dim1_interval_hist_table = interval_hist_table(dim1, 10)
    add_table(doc, dim1_interval_hist_table)
    add_paragraph(doc, intervals_hist_table_info)
    add_paragraph(doc, "График равноинтервальной гистограммы относительных частот:")
    dim1_interval_hist_plot = interval_hist_plot(dim1, 10)
    add_picture(doc, dim1_interval_hist_plot)
    add_paragraph(doc, "1.4. Равновероятностная гистограмма относительных частот\n"+probability_hist_info)
    add_paragraph(doc, "Расчётная таблица:")
    dim1_probability_hist_table = probability_hist_table(dim1, 10)
    add_table(doc, dim1_probability_hist_table)
    add_paragraph(doc, "\nГрафик равновероятностной гистограммы относительных частот:")
    dim1_probability_hist_plot = probability_hist_plot(dim1, 10)
    add_picture(doc, dim1_probability_hist_plot)
    dim1_avg = round(avg(dim1), 2)
    add_paragraph(doc, "1.5. Точечная оценка математического ожидания\n"+math_wait_info+f"\nДля данного вариационного ряда:\n\t\t\t\t\tmₓ = {dim1_avg}.")
    dim1_variance = round(variance(dim1, fixed=True), 2)
    add_paragraph(doc, "1.6. Точечная оценка дисперсии\n"+variance_info+f"\nДля данного вариационного ряда:\n\t\t\t\t\ts² = {dim1_variance}.")
    dim1_std = round(std(dim1, fixed=True), 2)
    dim1_delta = round(delta(dim1, 1.984), 2)
    dim1_avg_interval = avg_interval(dim1, 1.984)
    dim1_avg_interval = tuple(round(x, 2) for x in dim1_avg_interval)
    add_paragraph(doc, "1.7. Оценка доверительного интервала генеральной средней (γ = 0.95)\n"+math_wait_interval_info+f"\nДля данной выборки:\n\t\t\t\t\ts = {dim1_std}\n\t\t\t\tk = 99, α = 0.05, tᵧ = 1.984,\n\t\t\t\t\tδ = {dim1_delta},\n\t\t\t\txᵣ ∈ {dim1_avg_interval}.")
    dim1_varience_interval = varience_interval(dim1, 128.42, 73.361)
    dim1_varience_interval = tuple(round(x, 2) for x in dim1_varience_interval)
    add_paragraph(doc, "1.8. Оценка доверительного интервала генеральной дисперсии (γ = 0.95)\n"+varience_interval_info+f"\nДля данной выборки:\n\t\t\t\ta = 0.025, b = 0.975, k = 99,\n\t\t\t\t(χᵃₖ)² = 128.42, (χᵇₖ)² = 73.361,\n\t\t\t\tσ² ∈ {dim1_varience_interval}.")
    add_paragraph(doc, "1.9. Гипотеза о законе распределения случайной величины по критерию согласия Пирсона (α = 0.05)\n"+hypothesis_info+f"\n\tH₀: генеральная совокупность распределена нормально:\n\t\tF(x) = F₀(x).\n\tH₁: генеральная совокупность не распределена нормально:\n\t\tF(x) ≠ F₀(x).")
    add_paragraph(doc, pirson_info)
    add_paragraph(doc, "Расчётная таблица:")
    dim1_hypothesis_table = hypothesis_table(dim1, 10)  
    add_table(doc, dim1_hypothesis_table)
    dim1_xi_quad = round(xi_quad(dim1, 10), 2)
    if dim1_xi_quad > 14.067:
        res = "Так как (χ²)' > χ², то гипотеза H₀ отвергается."
    else:
        res = "Так как (χ²)' < χ², то гипотеза H₀ принимается (нет оснований для отвержения)."
    add_paragraph(doc, f"\nОткуда (χ²)' = {dim1_xi_quad}.")
    add_paragraph(doc, xi_quad_info+f"\nДля данной выборки:\n\t\t\tk = 7, α = 0.05, χ² = 14.067.\n\t{res}")
    dim1_colmogor = round(colmogor(dim1), 2)
    if dim1_colmogor > 1.36:
        res = "Так как K' > K, то гипотеза H₀ отвергается."
    else:
        res = "Так как K' < K, то гипотеза H₀ принимается (нет оснований для отвержения)."
    add_paragraph(doc, "1.10. Гипотеза о законе распределения случайной величины по критерию Колмогорова (α = 0.05)\n"+colmogor_info+f"\nДля данной выборки:\n\t\t\t\t\tK' = {dim1_colmogor},\n\t\t\t\tα = 0.05, K = 1.36.\n\t{res}")
    add_paragraph(doc, "График нормальной функции распределения и эмпирической функции распределения для данной выборки:")
    dim1_norm_empiric_plot = norm_empiric_distrib(dim1)
    add_picture(doc, dim1_norm_empiric_plot)
    add_paragraph(doc, "\n2. Анализ двумерной выборки",
                  align=WD_ALIGN_PARAGRAPH.CENTER, pt=16)
    add_paragraph(doc, "2.1. Точечная оценка коэффициента корреляции\n"+correlation_info)
    add_paragraph(doc, "Расчётная таблица:")
    dim2_xy_table = xy_table(dim2)
    x, y = zip(*dim2)
    x = np.array(x)
    y = np.array(y)
    x_average = round(avg(x), 3)
    y_average = round(avg(y), 3)
    xy_average = round(avg(x*y), 3)
    x_std = round(std(x), 3)
    y_std = round(std(y), 3)
    r = round((xy_average-x_average*y_average)/(x_std*y_std), 2)
    add_table(doc, dim2_xy_table)
    add_paragraph(doc, f"\nОткуда\n\t\t\tmˣʸ = {xy_average}, mˣ = {x_average}, mʸ = {y_average},\n\t\t\t\tsˣ = {x_std}, sʸ = {y_std},\n\t\t\t\t\tR = {r}.")
    a, b = a_b(dim2, r, 1.96)
    a = round(a, 2)
    b = round(b, 2)
    r_interval = correlation_interval(x, r, 1.96)
    r_interval = tuple(round(xi, 2) for xi in r_interval)
    add_paragraph(doc, "2.2. Оценка доверительного интервала генерального коэффициента корреляции (γ = 0.95)\n"+interval_correlation_info+f"\nДля данной выборки:\n\t\t\targФ(zᵧ) = γ/2 = 0.475, zᵧ = 1.96,\n\t\t\t\ta = {a}, b = {b},\n\t\t\t\tRᵣ ∈ {r_interval}.")
    dim2_static_crit = round(static_crit(len(dim2), r), 2)
    if abs(dim2_static_crit) > 2.01:
        res = "Так как |T'| > T, то гипотеза H₀ отвергается."
    else:
        res = "Так как |T'| < T, то гипотеза H₀ принимается (нет оснований для отвержения)."
    add_paragraph(doc, "2.3. Гипотеза об отсутствии корреляционной зависимости (α = 0.05)\n"+corr_hypo_info+f"\nДля данной выборки:\n\t\t\t\t\tT' = {dim2_static_crit},\n\t\t\t\tk = 48, T = 2.01.\n\t{res}")
    dim2_lin_reg = N(lin_reg(dim2, r), 2)
    add_paragraph(doc, "2.4. Построение линейной регрессии и диаграммы рассеяния\n"+lin_reg_info+f"\nВыразим из уравнения y и получим:\n\ty = {dim2_lin_reg}.")
    add_paragraph(doc, "График линейной регрессии:")
    dim2_lin_reg_plot = lin_reg_plot(dim2, r)
    add_picture(doc, dim2_lin_reg_plot)
    add_paragraph(doc, hint)
    doc.save(filename)


def add_table(doc: Document, df: DataFrame) -> None:
    rows = df.shape[0] + 1
    cols = df.shape[1]
    doc.add_table(rows=rows, cols=cols).style = 'Table Grid'
    for i in range(rows):
        for j in range(cols):
            if i == 0:
                doc.tables[-1].cell(i, j).text = df.columns[j]
                doc.tables[-1].cell(i, j).paragraphs[0].runs[0].italic = True
            else:
                value = df.iloc[i-1, j]
                if np.isnan(value):
                    value = ''
                elif value % 1 == 0:
                    value = value.astype("int64")
                else:
                    value = round(value, 2)
                doc.tables[-1].cell(i, j).text = str(value)
            doc.tables[-1].cell(i,
                                j).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.tables[-1].cell(i, j).alignment = WD_ALIGN_VERTICAL.CENTER


def add_picture(doc: Document, buffer: str) -> None:
    doc.add_picture(buffer)
    p = doc.paragraphs[-1]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_paragraph(doc: Document, text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, pt=14, font="Times New Roman") -> None:
    p = doc.add_paragraph(text)
    p.alignment = align
    _set_pt(p, pt)
    _set_font(p, font)


def _set_pt(p, size: int) -> None:
    run = p.runs[0]
    font = run.font
    font.size = Pt(size)


def _set_font(p, font: str) -> None:
    run = p.runs[0]
    rfont = run.font
    rfont.name = font


if __name__ == "__main__":
    create_all("export/V9_402.docx")
